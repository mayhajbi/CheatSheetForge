"""
בדיקת קבלה לאדם C (פרק 12): sample_merged.json -> docx/pdf תקינים, כולל
בדיקת כיווניות בסיסית (bidi=1 קיים בפסקאות עבריות).

הרצה: pytest backend/export/tests/test_export.py
(דורש pydantic + python-docx מותקנים, וגם soffice זמין ב-PATH לבדיקת PDF)
"""

import json
import shutil
from pathlib import Path

import pytest
from docx import Document

from backend.export.docx_builder import build_docx
from backend.export.pdf_builder import build_pdf
from backend.schemas import MergedBank

FIXTURE_PATH = Path(__file__).resolve().parents[3] / "fixtures" / "sample_merged.json"


def load_fixture_bank() -> MergedBank:
    data = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))
    return MergedBank(**data)


def test_build_docx_creates_file(tmp_path):
    bank = load_fixture_bank()
    out = tmp_path / "cheatsheet.docx"
    build_docx(bank, out)

    assert out.exists()
    doc = Document(str(out))
    assert len(doc.paragraphs) > 0


def test_docx_contains_all_topics(tmp_path):
    bank = load_fixture_bank()
    out = tmp_path / "cheatsheet.docx"
    build_docx(bank, out)

    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for item in bank.items:
        assert item.topic in all_text


@pytest.mark.skipif(shutil.which("soffice") is None, reason="LibreOffice לא מותקן בסביבה זו")
def test_build_pdf_creates_file(tmp_path):
    bank = load_fixture_bank()
    out = tmp_path / "cheatsheet.pdf"
    build_pdf(bank, out)

    assert out.exists()
    assert out.stat().st_size > 0


def test_hebrew_paragraphs_are_rtl_and_code_lines_are_ltr(tmp_path):
    """נקודת הסיכון המרכזית בפרויקט (פרק 05): כיווניות נכונה בקובץ שנוצר.

    פסקאות עברית -> bidi=1 ויישור לימין; שורות תשובה/קוד באנגלית -> bidi=0.
    """

    bank = load_fixture_bank()
    out = tmp_path / "cheatsheet.docx"
    build_docx(bank, out)

    doc = Document(str(out))
    by_dir = {"rtl": [], "ltr": []}
    for p in doc.paragraphs:
        xml = p._p.xml
        if 'w:bidi w:val="1"' in xml:
            by_dir["rtl"].append(p.text)
        elif 'w:bidi w:val="0"' in xml:
            by_dir["ltr"].append(p.text)

    # הכותרת ושאלות העברית בכיוון RTL
    assert any("דף נוסחאות" in t for t in by_dir["rtl"])
    assert any("עקביות אחרי קריסה" in t for t in by_dir["rtl"])

    # שורת הנוסחה/קוד באנגלית בשורה נפרדת בכיוון LTR (ולא מעורבת בעברית)
    assert any("avg_access_time" in t for t in by_dir["ltr"])
    assert all("avg_access_time" not in t for t in by_dir["rtl"])
