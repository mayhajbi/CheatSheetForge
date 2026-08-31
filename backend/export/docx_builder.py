"""
בניית קובץ Word (docx) מהבנק המרוכז, עם תמיכת RTL מלאה.

מבוסס על PoC שאומת ידנית (ראו docs/rtl_poc_notes.md): python-docx עם bidi=1
ברמת הפסקה שומר על כיווניות תקינה כאשר טקסט עברי (RTL) ונוסחה/קוד באנגלית
(LTR) מופרדים לשורות/פסקאות נפרדות -- בהתאם לכלל בפרק 05 של ה-PRD.

עקרון קשיח: אין להמציא תוכן. אם representative/variant חסר answer_text,
הוא מוצג עם סימון "(לא נמצאה תשובה במקור)" ולא מושלם מהזיכרון הכללי.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from backend.schemas import MergedBank

TOPIC_HEADINGS_ORDER = {"closed": "שאלות סגורות", "open_calc": "שאלות פתוחות", "code": "שאלות מימוש קוד"}


def _set_paragraph_direction(paragraph, rtl: bool) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1" if rtl else "0")
    pPr.append(bidi)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    )


def _add_rtl_paragraph(doc: Document, text: str, bold: bool = False, size_pt: int | None = None):
    p = doc.add_paragraph()
    _set_paragraph_direction(p, rtl=True)
    run = p.add_run(text)
    run.bold = bold
    if size_pt:
        run.font.size = Pt(size_pt)
    return p


def _add_ltr_code_paragraph(doc: Document, text: str, mono_font: str = "Consolas"):
    p = doc.add_paragraph()
    _set_paragraph_direction(p, rtl=False)
    run = p.add_run(text)
    run.font.name = mono_font
    return p


def _configure_base_document(doc: Document, font_name: str, font_size_pt: int) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:cs"), font_name)


def _render_closed_item(doc: Document, item: dict) -> None:
    _add_rtl_paragraph(doc, item["question_text"])
    _add_rtl_paragraph(doc, f"תשובה נכונה: {item['correct_answer']}")
    if item.get("distractors"):
        _add_rtl_paragraph(doc, "תשובות מטעות: " + " / ".join(item["distractors"]))
    if item.get("sources"):
        _add_rtl_paragraph(doc, "מקור: " + ", ".join(item["sources"]))


def _render_open_calc_item(doc: Document, item: dict) -> None:
    rep = item["representative"]
    _add_rtl_paragraph(doc, rep["question_text"])
    answer = rep.get("answer_text") or "(לא נמצאה תשובה במקור)"
    _add_ltr_code_paragraph(doc, answer)
    _add_rtl_paragraph(doc, f"מקור: {rep['source_label']}")

    variants = item.get("variants") or []
    if variants:
        _add_rtl_paragraph(doc, f"וריאציות נוספות ({len(variants)}):")
        for v in variants:
            _add_rtl_paragraph(doc, v["question_text"])
            v_answer = v.get("answer_text") or "(לא נמצאה תשובה במקור)"
            _add_ltr_code_paragraph(doc, v_answer)
            _add_rtl_paragraph(doc, f"מקור: {v['source_label']}")


def _render_code_item(doc: Document, item: dict) -> None:
    sources = ", ".join(item.get("sources", []))
    if item.get("reference_only", True):
        _add_rtl_paragraph(doc, f"שאלת מימוש קוד -- הפניה בלבד למקור: {sources}")
    else:
        snippet = item.get("code_snippet") or "(קוד חסר במקור)"
        _add_ltr_code_paragraph(doc, snippet)
        _add_rtl_paragraph(doc, f"מקור: {sources}")


_RENDERERS = {
    "closed": _render_closed_item,
    "open_calc": _render_open_calc_item,
    "code": _render_code_item,
}


def build_docx(
    bank: MergedBank,
    output_path: Path,
    font_name: str = "Arial",
    font_size_pt: int = 11,
) -> Path:
    """בונה קובץ docx מהבנק המרוכז ושומר ל-output_path."""

    doc = Document()
    _configure_base_document(doc, font_name, font_size_pt)

    _add_rtl_paragraph(doc, f"דף נוסחאות -- {bank.course}", bold=True, size_pt=16)

    items_by_type: dict[str, list[dict]] = {"closed": [], "open_calc": [], "code": []}
    for item in bank.items:
        items_by_type.setdefault(item["type"], []).append(item)

    for qtype in ("closed", "open_calc", "code"):
        items = items_by_type.get(qtype, [])
        if not items:
            continue
        _add_rtl_paragraph(doc, TOPIC_HEADINGS_ORDER[qtype], bold=True, size_pt=13)

        by_topic: dict[str, list[dict]] = {}
        for item in items:
            by_topic.setdefault(item["topic"], []).append(item)

        for topic, topic_items in by_topic.items():
            _add_rtl_paragraph(doc, topic, bold=True, size_pt=12)
            for item in topic_items:
                _RENDERERS[qtype](doc, item)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
